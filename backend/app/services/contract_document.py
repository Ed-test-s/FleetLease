"""Generate DOCX contracts from Word templates by replacing [placeholder] markers."""

import re
from datetime import timedelta
from io import BytesIO
from pathlib import Path

from docx import Document
from lxml import etree

from app.core.config import settings
from app.services.storage import storage_service

TEMPLATES_DIR = Path(__file__).resolve().parent.parent.parent / "contracts_examples"
PSA_TEMPLATE = TEMPLATES_DIR / "шаблон_договор_купли_продажи.docx"
LA_TEMPLATE = TEMPLATES_DIR / "шаблон_договор_лизинга.docx"


# ---------------------------------------------------------------------------
# Number-to-words
# ---------------------------------------------------------------------------

def _num_to_words(n: float, currency: str = "BYN") -> str:
    n = round(n, 2)
    int_part = int(n)
    frac = round((n - int_part) * 100)

    if int_part == 0:
        return f"ноль {_currency_word(currency)} {frac:02d} коп."

    groups: list[int] = []
    remaining = int_part
    while remaining > 0:
        groups.append(remaining % 1000)
        remaining //= 1000

    ones = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    ones_f = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    teens = ["десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать",
             "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
    tens = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят",
            "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
    hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот",
                "шестьсот", "семьсот", "восемьсот", "девятьсот"]

    def _group_words(g: int, feminine: bool = False) -> str:
        parts: list[str] = []
        h, t, o = g // 100, (g % 100) // 10, g % 10
        if h:
            parts.append(hundreds[h])
        if t == 1:
            parts.append(teens[o])
        else:
            if t:
                parts.append(tens[t])
            if o:
                parts.append((ones_f if feminine else ones)[o])
        return " ".join(parts)

    scale_names = [
        ("", "", "", False),
        ("тысяча", "тысячи", "тысяч", True),
        ("миллион", "миллиона", "миллионов", False),
        ("миллиард", "миллиарда", "миллиардов", False),
    ]
    result_parts: list[str] = []
    for i, g in enumerate(groups):
        if g == 0:
            continue
        sn = scale_names[i] if i < len(scale_names) else ("", "", "", False)
        w = _group_words(g, sn[3])
        o, t = g % 10, (g % 100) // 10
        if sn[0]:
            if t == 1:
                suffix = sn[2]
            elif o == 1:
                suffix = sn[0]
            elif 2 <= o <= 4:
                suffix = sn[1]
            else:
                suffix = sn[2]
            w = f"{w} {suffix}"
        result_parts.append(w)

    result_parts.reverse()
    text = " ".join(result_parts).strip()
    return f"{text} {_currency_word(currency)} {frac:02d} коп."


def _currency_word(code: str) -> str:
    return {"BYN": "руб.", "USD": "долл. США", "EUR": "евро", "RUB": "росс. руб."}.get(code, "руб.")


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def _format_date_text(d) -> str:
    """'1 января 2026 г.'"""
    if d is None:
        return "___________"
    if hasattr(d, "strftime"):
        months = ["", "января", "февраля", "марта", "апреля", "мая", "июня",
                  "июля", "августа", "сентября", "октября", "ноября", "декабря"]
        return f"{d.day} {months[d.month]} {d.year} г."
    return str(d)


def _format_date_dots(d) -> str:
    """'dd.mm.yyyy'"""
    if d is None:
        return "___________"
    if hasattr(d, "strftime"):
        return d.strftime("%d.%m.%Y")
    return str(d)


def _safe(val, default="___________") -> str:
    if val is None or val == "":
        return default
    return str(val)


# ---------------------------------------------------------------------------
# User data extractors
# ---------------------------------------------------------------------------

def _user_name(u) -> str:
    if not u:
        return ""
    if u.company:
        return u.company.company_name or ""
    if u.entrepreneur:
        return u.entrepreneur.full_name or ""
    if u.individual:
        return u.individual.full_name or ""
    return ""


def _legal_form(u) -> str:
    if u and u.company:
        return u.company.legal_form or ""
    return ""


def _legal_form_and_name(u) -> str:
    if not u:
        return ""
    if u.company:
        lf = u.company.legal_form or ""
        name = u.company.company_name or ""
        return f'{lf} "{name}"' if lf else f'"{name}"'
    if u.entrepreneur:
        return f'ИП "{u.entrepreneur.full_name or ""}"'
    if u.individual:
        return f'Физическое лицо "{u.individual.full_name or ""}"'
    return ""


def _director_or_name(u) -> str:
    if not u:
        return ""
    if u.company:
        return u.company.director_name or ""
    if u.entrepreneur:
        return u.entrepreneur.full_name or ""
    if u.individual:
        return u.individual.full_name or ""
    return ""


def _unp(u) -> str:
    if not u:
        return ""
    if u.company:
        return u.company.unp or ""
    if u.entrepreneur:
        return u.entrepreneur.unp or ""
    return ""


def _unp_or_dash(u) -> str:
    v = _unp(u)
    return v if v else "—"


def _legal_address(u) -> str:
    if not u:
        return ""
    if u.company:
        return u.company.legal_address or ""
    if u.entrepreneur:
        return u.entrepreneur.legal_address or ""
    if u.individual:
        return u.individual.registration_address or ""
    return ""


def _postal_address(u) -> str:
    if not u:
        return ""
    if u.company:
        return u.company.postal_address or u.company.legal_address or ""
    if u.entrepreneur:
        return u.entrepreneur.postal_address or u.entrepreneur.legal_address or ""
    if u.individual:
        return u.individual.registration_address or ""
    return ""


def _contact_phone(u) -> str:
    if not u:
        return ""
    return next((c.value for c in (u.contacts or []) if c.type.value == "phone"), "")


def _contact_email(u) -> str:
    if not u:
        return ""
    return next((c.value for c in (u.contacts or []) if c.type.value == "email"), "")


def _first_bank(u):
    return u.bank_accounts[0] if u and u.bank_accounts else None


def _passport_id(u) -> str:
    if u and u.individual:
        return u.individual.passport_id or ""
    return ""


def _to_initials(full_name: str) -> str:
    """'Иванов Иван Сергеевич' → 'Иванов И. С.'"""
    parts = full_name.strip().split()
    if not parts:
        return ""
    if len(parts) == 1:
        return parts[0]
    surname = parts[0]
    initials = " ".join(f"{p[0]}." for p in parts[1:] if p)
    return f"{surname} {initials}"


def _director_initials(u) -> str:
    """Director/name in 'Фамилия И. О.' format."""
    return _to_initials(_director_or_name(u))


# ---------------------------------------------------------------------------
# DOCX replacement engine
# ---------------------------------------------------------------------------

def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    full_text = paragraph.text
    if "[" not in full_text:
        return
    new_text = full_text
    for key, value in replacements.items():
        new_text = new_text.replace(f"[{key}]", str(value))
    if new_text == full_text:
        return
    if not paragraph.runs:
        return
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = new_text


def _replace_in_table(table, replacements: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_in_paragraph(p, replacements)


def _replace_in_doc(doc: Document, replacements: dict[str, str]) -> None:
    for p in doc.paragraphs:
        _replace_in_paragraph(p, replacements)
    for table in doc.tables:
        _replace_in_table(table, replacements)
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header and header.is_linked_to_previous is False:
                for p in header.paragraphs:
                    _replace_in_paragraph(p, replacements)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer and footer.is_linked_to_previous is False:
                for p in footer.paragraphs:
                    _replace_in_paragraph(p, replacements)


# ---------------------------------------------------------------------------
# Contextual text builders (for [1]…[9] and [101],[102])
# ---------------------------------------------------------------------------

def _build_seller_intro(supplier) -> str:
    """[1] — intro line for Продавец in PSA header."""
    if not supplier:
        return "___________"
    if supplier.company:
        lf = supplier.company.legal_form or ""
        cn = supplier.company.company_name or ""
        dn = supplier.company.director_name or ""
        return (f'{lf} "{cn}", именуемое в дальнейшем «Продавец», '
                f'в лице Директора {dn}, действующего на основании Устава, с одной стороны')
    if supplier.entrepreneur:
        fn = supplier.entrepreneur.full_name or ""
        return f'ИП "{fn}", именуемое в дальнейшем «Продавец», с одной стороны'
    if supplier.individual:
        fn = supplier.individual.full_name or ""
        return f'Физическое лицо "{fn}", именуемое в дальнейшем «Продавец», с одной стороны'
    return "___________"


def _build_buyer_intro(lessor) -> str:
    """[2] — intro line for Покупатель (always company)."""
    if not lessor or not lessor.company:
        return "___________"
    lf = lessor.company.legal_form or ""
    cn = lessor.company.company_name or ""
    dn = lessor.company.director_name or ""
    return (f'{lf} "{cn}", именуемое в дальнейшем «Покупатель», '
            f'в лице директора {dn}')


def _build_condition_full(vehicle) -> str:
    """[3] — 'новый' or 'бывший в эксплуатации'."""
    if vehicle and vehicle.condition == "new":
        return "новый"
    return "бывший в эксплуатации"


def _build_condition_short(vehicle) -> str:
    """[4] — 'Новый' or 'Б/у'."""
    if vehicle and vehicle.condition == "new":
        return "Новый"
    return "Б/у"


def _build_lessee_for_psa(lessee) -> str:
    """[6] — lessee mention in PSA purpose paragraph."""
    if not lessee:
        return "___________"
    if lessee.company:
        lf = lessee.company.legal_form or ""
        cn = lessee.company.company_name or ""
        return f'{lf} "{cn}"'
    if lessee.entrepreneur:
        return f'ИП {lessee.entrepreneur.full_name or ""}'
    if lessee.individual:
        return f'физическому лицу {lessee.individual.full_name or ""}'
    return "___________"


def _build_lessor_name_block(lessor) -> str:
    """[7] — Lessor org form + company name for rekvizity table."""
    if not lessor or not lessor.company:
        return "___________"
    lf = lessor.company.legal_form or ""
    cn = lessor.company.company_name or ""
    return f'{lf} "{cn}"'


def _build_supplier_name_block(supplier) -> str:
    """[8] — Supplier name block in rekvizity table."""
    if not supplier:
        return "___________"
    if supplier.company:
        lf = supplier.company.legal_form or ""
        cn = supplier.company.company_name or ""
        return f'{lf} "{cn}"'
    if supplier.entrepreneur:
        return f'ИП "{supplier.entrepreneur.full_name or ""}"'
    if supplier.individual:
        return f'Физическое лицо {supplier.individual.full_name or ""}'
    return "___________"


def _build_lessee_agreement_line(lessee) -> str:
    """[9] — 'Согласовано с Лизингополучателем …' name block."""
    if not lessee:
        return "___________"
    if lessee.company:
        lf = lessee.company.legal_form or ""
        cn = lessee.company.company_name or ""
        return f'{lf} "{cn}"'
    if lessee.entrepreneur:
        return f'ИП {lessee.entrepreneur.full_name or ""}'
    if lessee.individual:
        return f'физическим лицом "{lessee.individual.full_name or ""}"'
    return "___________"


def _build_la_lessee_header(lessee) -> str:
    """[101] — Лизингополучатель header in LA."""
    if not lessee:
        return "___________"
    if lessee.company:
        lf = lessee.company.legal_form or ""
        cn = lessee.company.company_name or ""
        dn = lessee.company.director_name or ""
        return (f'{lf} "{cn}" (Лизингополучатель) в лице директора {dn} '
                f'действующего на основании Устава, с другой стороны')
    if lessee.entrepreneur:
        fn = lessee.entrepreneur.full_name or ""
        return f'ИП {fn} (Лизингополучатель) с другой стороны'
    if lessee.individual:
        fn = lessee.individual.full_name or ""
        return f'Физическое лицо {fn} (Лизингополучатель) с другой стороны'
    return "___________"


def _build_la_supplier_line(supplier) -> str:
    """[102] — supplier line in LA predmet section."""
    if not supplier:
        return "___________"
    if supplier.company:
        lf = supplier.company.legal_form or ""
        cn = supplier.company.company_name or ""
        unp = supplier.company.unp or ""
        return f'{lf} "{cn}", УНП – {unp}'
    if supplier.entrepreneur:
        fn = supplier.entrepreneur.full_name or ""
        unp = supplier.entrepreneur.unp or ""
        return f'ИП "{fn}", УНП – {unp}'
    if supplier.individual:
        fn = supplier.individual.full_name or ""
        pid = supplier.individual.passport_id or ""
        return f'Физическое лицо "{fn}", идентификационный номер паспорта – {pid}'
    return "___________"


def _build_la_lessee_name_block(lessee) -> str:
    """[103] — Лизингополучатель name block (no 'в лице директора…')."""
    if not lessee:
        return "___________"
    if lessee.company:
        lf = lessee.company.legal_form or ""
        cn = lessee.company.company_name or ""
        return f'{lf} "{cn}"'
    if lessee.entrepreneur:
        return f'ИП {lessee.entrepreneur.full_name or ""}'
    if lessee.individual:
        return f'Физическое лицо {lessee.individual.full_name or ""}'
    return "___________"


def _build_la_lessee_info_name(lessee) -> str:
    """Org-form + name for ИНФОРМАЦИЯ О ЛИЗИНГОПОЛУЧАТЕЛЕ block."""
    if not lessee:
        return "___________"
    if lessee.company:
        lf = lessee.company.legal_form or ""
        cn = lessee.company.company_name or ""
        return f'{lf} "{cn}"'
    if lessee.entrepreneur:
        return f'ИП {lessee.entrepreneur.full_name or ""}'
    if lessee.individual:
        return f'Физическое лицо {lessee.individual.full_name or ""}'
    return "___________"


# ---------------------------------------------------------------------------
# PSA document generation
# ---------------------------------------------------------------------------

def generate_psa_document(contract, vehicle, supplier_user, lessor_user, lessee_user) -> BytesIO:
    doc = Document(str(PSA_TEMPLATE))

    vat_rate = contract.vat_rate or 20
    currency = contract.currency or "BYN"
    price = vehicle.price if vehicle else 0
    quantity = contract.quantity or 1
    total_price = price * quantity
    vat_amount = round(total_price * vat_rate / 100, 2)
    price_without_vat = round(total_price - vat_amount, 2)
    half_price = round(total_price / 2, 2)
    three_days = (contract.signing_date + timedelta(days=3)) if contract.signing_date else None

    condition_raw = vehicle.condition if vehicle else "used"
    condition_full_new = ("Товар является новым. В связи с этим Продавец предоставляет "
                          "гарантию на передаваемый товар согласно условиям.")
    condition_full_used = ("Товар является бывшим в эксплуатации и передается Покупателю в том состоянии, "
                           "в котором он фактически находится на момент его передачи. В связи с этим Продавец "
                           "не предоставляет гарантию на передаваемый товар, обмену и возврату товар не подлежит.")
    warranty_new = "Продавец несет ответственность за скрытые недостатки товара."
    warranty_used = ("Продавец не несет ответственности за скрытые недостатки товара "
                     "в связи с тем, что товар является бывшим в эксплуатации.")

    supplier_bank = _first_bank(supplier_user)
    lessor_bank = _first_bank(lessor_user)

    repl: dict[str, str] = {
        # Contextual blocks [1]…[9]
        "1": _build_seller_intro(supplier_user),
        "2": _build_buyer_intro(lessor_user),
        "3": _build_condition_full(vehicle),
        "4": _build_condition_short(vehicle),
        "5": _format_date_dots(contract.tech_passport_date),
        "6": _build_lessee_for_psa(lessee_user),
        "7": _build_lessor_name_block(lessor_user),
        "8": _build_supplier_name_block(supplier_user),
        "9": _build_lessee_agreement_line(lessee_user),

        # Contract fields
        "номер ДКП": _safe(contract.contract_number),
        "год подписания": str(contract.signing_date.year) if contract.signing_date else "____",
        "город подписания": _safe(contract.signing_city),
        "дата подписания": _format_date_dots(contract.signing_date),

        # Vehicle
        "тип ТС": _safe(vehicle.vehicle_type if vehicle else None, "транспортное средство"),
        "наименование полностью": _safe(vehicle.name if vehicle else None),
        "год выпуска": _safe(vehicle.release_year if vehicle else None),
        "VIN-номер": _safe(vehicle.vin if vehicle else None),
        "№техпаспорта": _safe(contract.tech_passport_number),
        "кол-во": str(quantity),

        # Financial
        "цена в объявлении": f"{total_price:.2f}",
        "валюта": _safe(currency, "BYN"),
        "число НДС": str(int(vat_rate)),
        "число НДС на момент подписания": str(int(vat_rate)),
        "сумма НДС": f"{vat_amount:.2f}",
        "цена в объявлении – сумма НДС": f"{price_without_vat:.2f}",
        "сумма словами": _num_to_words(total_price, currency),
        "сумма НДС словами": _num_to_words(vat_amount, currency),
        "цена из объявления словами": _num_to_words(total_price, currency),
        "половина от полной цены в объявлении": f"{half_price:.2f}",
        "половина от полной цены словами": _num_to_words(half_price, currency),
        "способ оплаты": "безналичного перечисления денежных средств",
        "дата, спустя три дня со дня формирования договора": _format_date_dots(three_days),

        # Condition/warranty long strings (exact match from template)
        ("Товар является бывшим в эксплуатации и передается Покупателю в том состоянии, "
         "в котором он фактически находится на момент его передачи. В связи с этим Продавец "
         "не предоставляет гарантию на передаваемый товар, обмену и возврату товар не подлежит. "
         "ИЛИ: Товар является новым. В связи с этим Продавец предоставляет гарантию на передаваемый "
         "товар согласно условиям."): (
            condition_full_new if condition_raw == "new" else condition_full_used),
        ("Продавец не несет ответственности за скрытые недостатки товара в связи с тем, "
         "что товар является бывшим в эксплуатации. ИЛИ: Продавец несет ответственность "
         "за скрытые недостатки товара."): (
            warranty_new if condition_raw == "new" else warranty_used),

        # Supplier data
        "юридический адрес поставщика": _legal_address(supplier_user),
        "УНП поставщика": _unp_or_dash(supplier_user),
        "IBAN поставщика": _safe(supplier_bank.iban if supplier_bank else None),
        "название банка поставщика": _safe(supplier_bank.bank_name if supplier_bank else None),
        "SWIFT поставщика": _safe(supplier_bank.swift if supplier_bank else None, ""),
        "BIC поставщика": _safe(supplier_bank.bic if supplier_bank else None, ""),
        "ФИО директора поставщика": _director_initials(supplier_user),
        "номер телефона поставщика": _contact_phone(supplier_user),
        "электронная почта поставщика": _contact_email(supplier_user),

        # Lessor data
        "УНП лизингодателя": _unp(lessor_user),
        "юридический адрес лизингодателя": _legal_address(lessor_user),
        "IBAN лизингодателя": _safe(lessor_bank.iban if lessor_bank else None),
        "название банка лизингодателя": _safe(lessor_bank.bank_name if lessor_bank else None),
        "SWIFT лизингодателя": _safe(lessor_bank.swift if lessor_bank else None, ""),
        "BIC лизингодателя": _safe(lessor_bank.bic if lessor_bank else None, ""),
        "ФИО директора лизингодателя": _director_initials(lessor_user),
        "номер телефона лизингодателя": _contact_phone(lessor_user),
        "электронная почта лизингодателя": _contact_email(lessor_user),

        # Lessee
        "ФИО директора лизингополучателя": _director_initials(lessee_user),

        # Conditional labels (#58, #59)
        "Директор": "Директор" if (lessee_user and lessee_user.company) else "",
        "Директор1": "Директор" if (supplier_user and supplier_user.company) else "",
    }

    _replace_in_doc(doc, repl)
    _protect_document(doc)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# LA document generation
# ---------------------------------------------------------------------------

def generate_la_document(contract, vehicle, lessor_user, lessee_user, supplier_user=None) -> BytesIO:
    doc = Document(str(LA_TEMPLATE))

    vat_rate = contract.vat_rate or 20
    currency = contract.currency or "BYN"
    price = vehicle.price if vehicle else 0
    total = contract.total_amount or 0
    vat_on_total = round(total * vat_rate / 100, 2)
    vat_on_price = round(price * vat_rate / 100, 2)
    price_without_vat = round(price - vat_on_price, 2)

    condition = "Новый" if (vehicle and vehicle.condition == "new") else "Б/у"

    lessor_bank = _first_bank(lessor_user)
    lessee_bank = _first_bank(lessee_user)

    repl: dict[str, str] = {
        # Contextual blocks
        "101": _build_la_lessee_header(lessee_user),
        "102": _build_la_supplier_line(supplier_user),
        "103": _build_la_lessee_name_block(lessee_user),

        # Contract
        "номер договора лизинга": _safe(contract.contract_number),
        "номер-договора": _safe(contract.contract_number),
        "город подписания": _safe(contract.signing_city),
        "дата подписания": _format_date_dots(contract.signing_date),
        "дата окончания срока": _format_date_text(contract.end_date),
        "валюта": _safe(currency, "BYN"),

        # Lessor header (always company)
        "Организационно правовая форма лизингодателя": _legal_form(lessor_user),
        "Название компании лизингодателя": _user_name(lessor_user),
        "ФИО директора": _director_or_name(lessor_user),
        "ФИО директора лизингодателя": _director_or_name(lessor_user),
        "УНП лизингодателя": _unp(lessor_user),
        "юридический адрес лизингодателя": _legal_address(lessor_user),
        "почтовый адрес лизингодателя": _postal_address(lessor_user),
        "номера телефонов лизингодателя": _contact_phone(lessor_user),
        "электронная почта лизингодателя": _contact_email(lessor_user),
        "IBAN лиздат": _safe(lessor_bank.iban if lessor_bank else None),
        "название банка лиздат": _safe(lessor_bank.bank_name if lessor_bank else None),
        "BIC лиздат": _safe(lessor_bank.bic if lessor_bank else None, ""),
        "SWIFT лиздат": _safe(lessor_bank.swift if lessor_bank else None, ""),

        # Lessee info block (adapt org form for IE/individual)
        "Организационно правовая форма лизингополучателя": (
            _legal_form(lessee_user) if lessee_user and lessee_user.company
            else ("ИП" if lessee_user and lessee_user.entrepreneur
                  else ("Физическое лицо" if lessee_user and lessee_user.individual else ""))),
        "Название компании лизингополучателя": _user_name(lessee_user),
        "УНП лизингополучателя": _unp_or_dash(lessee_user),
        "юридический адрес лизингополучателя": _legal_address(lessee_user),
        "почтовый адрес лизингополучателя": _postal_address(lessee_user),
        "номера телефонов лизингополучателя": _contact_phone(lessee_user),
        "электронная почта лизингополучателя": _contact_email(lessee_user),
        "IBAN лизполуч": _safe(lessee_bank.iban if lessee_bank else None),
        "название банка лизполуч": _safe(lessee_bank.bank_name if lessee_bank else None),
        "BIC лизполуч": _safe(lessee_bank.bic if lessee_bank else None, ""),
        "SWIFT лизполуч": _safe(lessee_bank.swift if lessee_bank else None, ""),

        # Vehicle
        "состояние": condition,
        "тип ТС": _safe(vehicle.vehicle_type if vehicle else None, "транспортное средство"),
        "название полностью": _safe(vehicle.name if vehicle else None),
        "год выпуска": _safe(vehicle.release_year if vehicle else None),
        "VIN номер": _safe(vehicle.vin if vehicle else None),
        "кол-во": str(contract.quantity or 1),

        # Financial
        "цена в объявлении": f"{price:.2f}",
        "цена в объявлении – сумма НДС": f"{price_without_vat:.2f}",
        "сумма НДС": f"{vat_on_price:.2f}",
        "текущий НДС": str(int(vat_rate)),
        "число НДС на момент подписания": str(int(vat_rate)),
        "сумма договора лизинга": f"{total:.2f}",
        "число": f"{total:.2f}",
        "число словами": _num_to_words(total, currency),
        "число суммы договора лизинга словами": _num_to_words(total, currency),
        "сумма НДС от договора лизинга": f"{vat_on_total:.2f}",
        "сумма НДС от договора лизинга словами": _num_to_words(vat_on_total, currency),

        # Signatures
        "ФИО лизингополучателя": _director_initials(lessee_user),
        "ФИО директора лизингодателя": _director_initials(lessor_user),

        # Bottom table: Лизингодатель
        "Организационно правовая форма лиздат": _legal_form(lessor_user),
        "Название компании": _user_name(lessor_user),
        "Название компании лиздат": _user_name(lessor_user),
        "ФИО директора лиздат": _director_initials(lessor_user),

        # Bottom table: Лизингополучатель
        "ФИО директора лизпоч": _director_initials(lessee_user),
    }

    _replace_in_doc(doc, repl)
    _protect_document(doc)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Document protection (read-only)
# ---------------------------------------------------------------------------

_WORDPROCESSINGML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _protect_document(doc: Document) -> None:
    """Set document protection to readOnly — prevents editing in Word."""
    settings_elem = doc.settings.element
    ns = _WORDPROCESSINGML_NS
    protection = settings_elem.find(f"{{{ns}}}documentProtection")
    if protection is None:
        protection = etree.SubElement(settings_elem, f"{{{ns}}}documentProtection")
    protection.set(f"{{{ns}}}edit", "readOnly")
    protection.set(f"{{{ns}}}enforcement", "1")


# ---------------------------------------------------------------------------
# Upload
# ---------------------------------------------------------------------------

def contract_docx_filename(prefix: str, contract_number: str | None, *, fallback_suffix: str | int | None = None) -> str:
    """Имя .docx для объекта в хранилище и для скачивания (последний сегмент URL)."""
    raw = (contract_number or "").strip()
    safe = re.sub(r'[\\/:*?"<>|\r\n]+', "_", raw)
    safe = safe.strip(" .")
    if not safe:
        safe = str(fallback_suffix) if fallback_suffix is not None else "без_номера"
    return f"{prefix}_{safe}.docx"


def upload_contract_document(buf: BytesIO, folder: str, filename: str) -> str:
    """Upload a protected DOCX to MinIO. Returns the URL."""
    data = buf.read()
    object_name = f"{folder}/{filename}"
    storage_service.client.put_object(
        settings.MINIO_BUCKET,
        object_name,
        BytesIO(data),
        length=len(data),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    return f"http://{settings.MINIO_EXTERNAL_ENDPOINT}/{settings.MINIO_BUCKET}/{object_name}"
